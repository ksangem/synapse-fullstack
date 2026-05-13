"""Build Synapse BRD v1.0 — Nalashaa Digital branding, ~30 pages.

Run:  python _build_brd.py
Output: Synapse_BRD_v1.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "Synapse_BRD_v1.docx"
DIAGRAMS = Path(__file__).parent / "diagrams"

NAVY = RGBColor(0x00, 0x20, 0x60)
H2_BLUE = RGBColor(0x00, 0xB0, 0xF0)
ORANGE = RGBColor(0xFF, 0x99, 0x33)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
ALTROW = "EBF3FB"
HDR_FILL = "002060"
BORDER_GREY = "BFBFBF"

doc = Document()

# ============== PAGE SETUP ==============
section = doc.sections[0]
section.page_height = Cm(27.94)
section.page_width = Cm(21.59)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Cm(0.97)
section.footer_distance = Cm(0.97)

# ============== STYLES ==============
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ============== HELPERS ==============

def set_cell_shading(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER_GREY, sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "002060")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = H2_BLUE
    return p


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = BLACK
    return p


def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.bold = bold
    r.font.italic = italic
    return p


def add_image(filename, caption=None, width_in=6.5):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run()
    img_path = DIAGRAMS / filename
    if img_path.exists():
        r.add_picture(str(img_path), width=Inches(width_in))
    else:
        r.add_text(f"[diagram missing: {filename}]")
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(12)
        cr = cp.add_run(caption)
        cr.font.name = "Calibri"; cr.font.size = Pt(9); cr.font.italic = True
        cr.font.color.rgb = GREY


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(it)
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def add_table(headers, rows, col_widths_in=None):
    """Standard branded table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # autofit off
    table.autofit = False
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HDR_FILL)
        set_cell_border(cell)
        cell.paragraphs[0].text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = WHITE
    # Data rows
    for ri, row in enumerate(rows, start=1):
        is_alt = (ri % 2 == 1)
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if is_alt:
                set_cell_shading(cell, ALTROW)
            set_cell_border(cell)
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(str(val))
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            r.font.color.rgb = BLACK
    return table


# ============== HEADER & FOOTER ==============

def setup_header_footer():
    sec = doc.sections[0]
    # Header
    header = sec.header
    hp = header.paragraphs[0]
    hp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("Nalashaa | Synapse Integration Platform — BRD")
    hr.font.name = "Calibri"
    hr.font.size = Pt(9)
    hr.font.color.rgb = GREY
    # Header bottom border
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "A6A6A6")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Footer — 3-col borderless table
    footer = sec.footer
    # clear default paragraph
    fp = footer.paragraphs[0]
    fp.text = ""
    ft = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    ft.autofit = False
    widths = [Inches(2.0), Inches(2.5), Inches(2.0)]
    for i, w in enumerate(widths):
        ft.columns[i].cells[0].width = w
    # Make borderless
    tbl = ft._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Left cell — logo placeholder text (logo path not accessible in this env)
    lc = ft.rows[0].cells[0]
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run("[ND Logo]")
    lr.font.name = "Calibri"
    lr.font.size = Pt(8)
    lr.font.italic = True
    lr.font.color.rgb = GREY

    # Centre cell — confidentiality
    cc = ft.rows[0].cells[1]
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = cc.paragraphs[0]
    cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run("© 2026 Nalashaa Solutions India Pvt. Ltd. | Confidential — Internal Use Only")
    cr.font.name = "Calibri"
    cr.font.size = Pt(9)
    cr.font.color.rgb = GREY

    # Right cell — Page X of Y via field codes
    rc = ft.rows[0].cells[2]
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]
    rp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pre = rp.add_run("Page ")
    pre.font.name = "Calibri"; pre.font.size = Pt(9); pre.font.color.rgb = GREY

    def add_field(paragraph, instr_text):
        run1 = paragraph.add_run()
        run1.font.name = "Calibri"; run1.font.size = Pt(9); run1.font.color.rgb = GREY
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        run1._r.append(fld_begin)
        run2 = paragraph.add_run()
        run2.font.name = "Calibri"; run2.font.size = Pt(9); run2.font.color.rgb = GREY
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = instr_text
        run2._r.append(instr)
        run3 = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fld_end)

    add_field(rp, " PAGE ")
    of = rp.add_run(" of ")
    of.font.name = "Calibri"; of.font.size = Pt(9); of.font.color.rgb = GREY
    add_field(rp, " NUMPAGES ")


# ============== COVER PAGE ==============

def cover():
    # spacer
    for _ in range(5):
        doc.add_paragraph()

    # BU
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("NALASHAA DIGITAL")
    r.font.name = "Calibri"; r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = ORANGE

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Synapse Integration Platform")
    r.font.name = "Calibri"; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY

    # Subtitle
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("Business Requirements Document & Developer Handover")
    r.font.name = "Calibri"; r.font.size = Pt(14); r.font.color.rgb = GREY

    # Meta block
    for line, bold in [
        ("Version 1.0", True),
        ("Date: 4 May 2026", False),
        ("Classification: Internal", False),
        ("BU: Nalashaa Digital (GLOBAL)", False),
        ("", False),
        ("Prepared for: Incoming Developer + Project Manager", False),
        ("Prepared by: Synapse Product Team", False),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(line)
        r.font.name = "Calibri"; r.font.size = Pt(11); r.font.bold = bold
        r.font.color.rgb = GREY if not bold else NAVY

    doc.add_page_break()


# ============== TOC ==============

def toc():
    add_h1("Table of Contents")
    p = doc.add_paragraph()
    r = p.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    r._r.append(fldChar)
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    r._r.append(fldEnd)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("[Right-click → Update Field in Word to populate the TOC]")
    r2.font.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GREY
    doc.add_page_break()


# ============== DOCUMENT CONTROL ==============

def doc_control():
    add_h1("1. Document Control")
    add_h2("1.1 Version History")
    add_table(
        ["Version", "Date", "Author", "Summary of Changes"],
        [
            ["0.1", "2026-04-22", "Product Team", "Internal status note (PRODUCT_STATUS.md) drafted."],
            ["0.5", "2026-04-30", "Product Team", "Charter v2.1 published (Synapse_Product_Charter_v2.pdf)."],
            ["1.0", "2026-05-04", "Product Team", "Initial BRD release for developer handover."],
        ],
        col_widths_in=[1.0, 1.1, 1.5, 2.9],
    )
    add_h2("1.2 Distribution List")
    add_table(
        ["Role", "Name", "Purpose"],
        [
            ["Product Owner", "Sandeep (Synapse initiator)", "Author of charter; ratifies scope changes"],
            ["Incoming Developer", "TBD", "Will pick up implementation of Phase 2+ features"],
            ["Project Manager", "TBD", "Will run the day-level project plan companion workbook"],
            ["Engineering Lead", "TBD", "Reviews architecture decisions, code merges"],
        ],
        col_widths_in=[1.6, 2.0, 2.9],
    )
    add_h2("1.3 Companion Documents")
    add_bullets([
        "Synapse_Product_Charter_v2.pdf — original product vision and feature definitions (read first if new to the product).",
        "Synapse_Project_Plan_v1.xlsx — day-level project plan and WBS for the project manager.",
        "PRODUCT_STATUS.md — in-repo developer-facing module status notes (kept up to date as code changes).",
        "CLAUDE.md / IMPLEMENTATION.md / SHAREPOINT_MODULE.md — in-repo developer playbooks for Claude Code (vibe coding) workflows.",
    ])
    doc.add_page_break()


# ============== EXECUTIVE SUMMARY ==============

def executive_summary():
    add_h1("2. Executive Summary")
    add_body(
        "Synapse is a self-hosted integration middleware platform that connects disparate systems — APIs, "
        "databases, file shares, SaaS applications — through a visual, drag-and-drop interface. The platform "
        "separates design-time activities (connector creation, entity modelling, transformation scripting) from "
        "runtime operations (connection setup, field mapping, deployment, monitoring), so non-technical "
        "Operators can build integrations in under 30 minutes without writing code."
    )
    add_body(
        "Today the platform connects Jira (source) to SharePoint (destination) end-to-end with credential "
        "encryption, scheduling, delta sync, and a 35-field mapper. The next phase expands it into a "
        "multi-source / multi-destination adapter platform supporting Dynamics 365, Jira, TARA, PostgreSQL, "
        "Excel, Keka as sources and TFS, SharePoint, Holiday Tracker, Ahrefs, GSC, Google Adwords as "
        "destinations — alongside a Connector Studio, Master Entity Catalog, AI Auto-Mapping, Trading "
        "Network Console, RBAC, Web Scraping, Alerting, and Help System."
    )
    add_h2("2.1 Why This Document Exists")
    add_body(
        "This BRD has a single purpose: equip an incoming developer (working with Claude Code / vibe coding) to "
        "pick up the codebase cold and deliver Phases 2, 3, and 4 of the charter. It captures three things in one "
        "place: (a) the product vision, (b) what is actually built today and how it is wired, and (c) every "
        "feature still to be built — with acceptance criteria, dependencies, and the phase it belongs in."
    )
    add_h2("2.2 30-Second Read")
    add_bullets([
        "Stack: React 19 + Vite (frontend), Express 5 + TypeScript + Drizzle (backend), PostgreSQL 16, Redis 7 + BullMQ (queues), Playwright (browser auth).",
        "What works end-to-end: Jira → SharePoint (35-field mapping, OAuth2 + API token, AES-256 credentials, cron schedules, delta sync, dedup).",
        "What is stubbed: 10 of 11 frontend pages use mock data; Studio is disabled; auth/RBAC/audit/alerts/multi-tenancy not enforced.",
        "Next milestones: wire the frontend to real APIs, ship Connector Studio, add new connectors (Dynamics 365, PostgreSQL, Excel, Keka, TARA), then TFS / Ahrefs / GSC / Holiday Tracker / Google Adwords destinations.",
        "Pace: aggressive day-level plan (in companion Excel) — vibe-coded by 1 developer + Claude.",
    ])
    doc.add_page_break()


# ============== VISION ==============

def vision():
    add_h1("3. Vision & Business Drivers")
    add_h2("3.1 Vision Statement")
    add_body(
        "Build connectors once, use them many times with zero code. Any team member — regardless of "
        "technical skill — should be able to connect two systems, map their data, test the pipeline, and deploy "
        "it to production within 30 minutes."
    )
    add_h2("3.2 Problems Solved")
    add_bullets([
        "Every integration today is custom-built from scratch, consuming significant engineering time and creating duplication of effort.",
        "Credentials are stored insecurely — hard-coded in scripts, shared via email, or scattered across teams.",
        "There is no centralised monitoring — integration failures are typically discovered by end users, not by the system.",
        "Integration logic lives in undocumented individual scripts with no handover support, making team transitions painful.",
        "Adding a new data source requires developer involvement every time, creating a scale bottleneck.",
    ])
    add_h2("3.3 Strategic Drivers")
    add_table(
        ["Driver", "Outcome"],
        [
            ["Reduce integration delivery time", "From 2–4 weeks per integration to under 30 minutes for an Operator using a pre-built connector."],
            ["Eliminate hard-coded credentials", "100% of credentials managed via encrypted vault or runtime prompt — auditable code base."],
            ["Reduce error detection latency", "From end-user reports to under 5 minutes via real-time monitoring and break-fix alerting."],
            ["Enable non-engineers", "Business analysts and team leads can self-serve integrations through the visual wizard."],
            ["Centralise governance", "Per-client usage tiers, RBAC, action audit trail, and credential lifecycle in one console."],
        ],
        col_widths_in=[2.3, 4.2],
    )
    doc.add_page_break()


# ============== PERSONAS ==============

def personas():
    add_h1("4. Personas & Roles")
    add_body("Four personas drive every workflow. Their access is governed by RBAC (see §7.8).")
    add_table(
        ["Persona", "Role", "Skill Level", "Primary Activities", "Frequency"],
        [
            ["Designer", "Integration Engineer / Developer", "High", "Creates connector templates, defines entities, writes transformation scripts, manages API registrations.", "Per new system onboarding"],
            ["Operator", "Business Analyst / Team Lead / PM", "Low to None", "Picks pre-built connectors, maps fields visually, deploys integrations, monitors health.", "Daily / Weekly"],
            ["Administrator", "IT Admin / Platform Owner", "Medium", "Manages client registrations, credentials, user access, platform health, approvals.", "As needed"],
            ["Consumer App", "External system calling Synapse API", "N/A (machine)", "Sends or receives data via registered API credentials.", "Continuous (automated)"],
        ],
        col_widths_in=[1.0, 1.6, 0.9, 2.2, 0.8],
    )
    doc.add_page_break()


# ============== CURRENT STATE ==============

def current_state():
    add_h1("5. Current State — What Is Built Today")
    add_body(
        "This section is the most important part of this document for an incoming developer. It is the ground "
        "truth of what runs today, sourced from a fresh codebase walkthrough on 4 May 2026 and reconciled "
        "against the in-repo PRODUCT_STATUS.md."
    )

    add_h2("5.1 Module Status Matrix")
    add_table(
        ["Module", "Status", "Notes"],
        [
            ["Jira Source — REST API (Red Gold)", "Built", "Basic Auth + JQL search, retries, custom-field parsing, normalizer."],
            ["Jira Source — Browser Automation (Flatiron)", "Built", "Playwright SSO/MFA, TOTP, 8h cookie session cache."],
            ["SharePoint Destination", "Built", "OAuth2 client credentials, Graph API site/list resolution, 35-field push, batched 20/req."],
            ["Field Mapping (Jira → SharePoint)", "Built", "12 standard + 23 derived fields (CycleTimeDays, IsOverdue, SprintNumber, etc.)."],
            ["Credential Encryption", "Built", "AES-256-GCM, expiry timestamp, store-only (no GET/decrypt endpoint yet)."],
            ["Integration CRUD", "Built", "Zod-validated, real DB queries, JSONB mappings + retry policy."],
            ["Run Tracking", "Built", "Status pending → running → success/error, record counts, error log JSONB."],
            ["Cron Scheduling", "Built", "BullMQ repeatable jobs via SchedulerService; PATCH /connected/:id/schedule."],
            ["Delta Sync", "Built", "Watermark-based, 3 modes (RESYNC, EXTEND_TO_TODAY, CUSTOM), terminal-status detection."],
            ["3-Layer Dedup", "Built", "pushLog → jiraItemCache → SharePoint direct query."],
            ["Connection Wizard", "Partial", "6-step UI shell; only Step 2 (Test Connection) is wired to real APIs."],
            ["Mapping Canvas", "Partial", "Pure-frontend UI with simulated AI auto-map; not persisted to backend."],
            ["Push Page", "Partial", "Form-based; talks to /api/sharepoint/push."],
            ["Connector Studio", "Stub", "Buttons disabled (StudioPage.jsx:86); no backend endpoints."],
            ["Health Dashboard / Registry / Monitor / Alerts / Vault / Catalog / Connected / Admin", "Stub UI", "Beautifully built UI shells, all on mock data."],
            ["RBAC + User Auth", "Missing", "users + role enum + organizations tables exist; never enforced."],
            ["Multi-Tenancy", "Missing", "organizations table + hardcoded UUID; no isolation."],
            ["Alerts System", "Missing", "DB table + queue declared; no creation/dispatch logic."],
            ["Audit Log", "Missing", "DB table exists; no writer."],
            ["Web Scraping (Apify / Playwright actor)", "Missing", "Charter Phase 2; not started."],
            ["Database / File Share / Message Queue connectors", "Missing", "Charter Phase 2; not started."],
            ["AI Auto-Mapping", "Missing", "Stub UI in Canvas; no LLM integration."],
            ["Help System", "Missing", "HelpPanel widget exists; no content engine."],
        ],
        col_widths_in=[2.6, 0.9, 3.0],
    )

    add_h2("5.2 Working End-to-End Journey: Jira → SharePoint")
    add_image("data_flow.png",
              caption="Figure 1 — Synapse current data flow: Jira extraction (REST or Playwright) → Normalizer → Mapper → Push with 3-layer dedup → SharePoint, plus persistence and observability.",
              width_in=6.5)
    add_bullets([
        "Operator hits POST /api/jira/test-connection (or /browser-auth for SSO clients) — backend validates credentials.",
        "Operator hits POST /api/integrations to register an integration with field mappings, schedule cron, retry policy.",
        "BullMQ scheduler enqueues a job per cron tick (or POST /api/integrations/:id/run for manual).",
        "integration-runner worker fetches Jira via Red Gold REST or Flatiron browser path, normalises to NormalizedJiraTicket.",
        "SharePointMapperService transforms 12 standard + 23 derived fields per the integration's fieldMappings JSONB.",
        "SharePointPushService deduplicates against pushLog → jiraItemCache → live SharePoint query, then creates / updates list items in batches of 20 via Microsoft Graph.",
        "Run row updated with recordsIn / recordsOut / status; sharepoint_push_runs row updated with progress; sync_state watermarks advanced.",
        "Operator polls GET /api/sharepoint/progress/:pushRunId or GET /api/runs to see results.",
    ])

    add_h2("5.3 Tech Stack Snapshot")
    add_table(
        ["Layer", "Technology", "Version", "Notes"],
        [
            ["Frontend framework", "React + React Router", "19 / 7", "JavaScript (no TypeScript), Vite 8 dev server."],
            ["UI styling", "Plain CSS + CSS variables", "—", "No MUI / Tailwind / Ant; light/dark via variables."],
            ["State management", "React Context API", "—", "Theme, Toast, DetailPane, Sidebar."],
            ["API client", "Fetch wrapper", "services/api.js", "Base URL http://localhost:4000; null on error → mock fallback."],
            ["Backend runtime", "Node.js + TypeScript", "20+ / 6.0", "CommonJS."],
            ["Backend framework", "Express", "5", "CORS + JSON body parser."],
            ["ORM / DB", "Drizzle / PostgreSQL", "0.45 / 16", "2 schemas (app, jira_data); 14 tables."],
            ["Queue", "BullMQ + Redis", "5.73 / 7", "3 active workers + 2 declared-but-empty (alert-dispatcher, credential-rotator)."],
            ["Browser automation", "Playwright", "1.59", "Jira SSO/MFA via TOTP (otplib 13.4)."],
            ["Encryption", "AES-256-GCM", "Node crypto", "CredentialService module."],
            ["Validation", "Zod", "4.3", "All POST endpoints validated."],
            ["Tests", "Vitest", "4.1", "10 e2e + unit test files."],
            ["Containerisation", "Docker Compose", "—", "Postgres + Redis services in docker-compose.yml."],
        ],
        col_widths_in=[1.4, 1.6, 1.0, 2.5],
    )

    add_h2("5.4 Data Model — All 14 Tables")
    add_table(
        ["Schema.Table", "Purpose", "Status"],
        [
            ["app.organizations", "Multi-tenant root", "Created, hardcoded UUID"],
            ["app.users", "User accounts (admin/designer/operator/viewer)", "Created, never read"],
            ["app.credentials", "Encrypted credential payloads", "Active"],
            ["app.connectors", "Connector templates (catalog)", "Created, never used"],
            ["app.integrations", "Integration instances (mappings, schedule, retry)", "Active"],
            ["app.runs", "Execution history with counts and error log", "Active"],
            ["app.run_messages", "Per-message tracking inside a run", "Created, never written"],
            ["app.alerts", "Break-fix alerts queue", "Created, never written"],
            ["app.audit_log", "Permission-sensitive action history", "Created, never written"],
            ["app.sharepoint_push_runs", "Per-push progress + counts + errors", "Active"],
            ["app.push_log", "Push history + dedup keys (cliendId, projectKey, date range)", "Active"],
            ["app.sync_state", "Delta-sync watermarks per integration", "Active"],
            ["app.jira_item_cache", "JiraKey ↔ SharePoint item ID + terminal-status flag", "Active"],
            ["jira_data.jira_tickets", "Raw + normalised Jira issues per run", "Active"],
        ],
        col_widths_in=[2.2, 3.0, 1.3],
    )
    add_body("Active: 8. Created-but-unused: 6. Wiring up the unused six is part of Phase 2 and Phase 3 work.")

    add_h2("5.5 API Surface — 8 Route Groups")
    add_table(
        ["Mount", "Routes", "Notes"],
        [
            ["/api/integrations", "GET, POST, GET :id, PATCH :id, DELETE :id, POST :id/run", "CRUD + manual execution trigger."],
            ["/api/runs", "GET, GET :id", "Run history."],
            ["/api/credentials", "POST", "Store-only — no decrypt endpoint yet."],
            ["/api/jira", "POST /test-connection, /browser-auth, GET /browser-auth/status, POST /fetch, GET /projects, POST /discover-entities, /discover-projects, /entity-fields", "Both Red Gold + Flatiron paths."],
            ["/api/sharepoint", "POST /test-connection, /list-fields, /push, GET /progress/:pushRunId", "Auth + push + progress polling."],
            ["/api/push", "POST /", "Direct push without integration."],
            ["/api/sync", "POST /:integrationId/start", "Trigger delta sync (RESYNC / EXTEND_TO_TODAY / CUSTOM)."],
            ["/api/connected", "GET /, PATCH /:integrationId/schedule", "Integration dashboard + cron management."],
        ],
        col_widths_in=[1.4, 3.0, 2.1],
    )
    add_body("Health probe: GET /health → { status: 'ok' }.")

    add_h2("5.6 Authentication & Security Posture Today")
    add_bullets([
        "Credential storage: AES-256-GCM in app.credentials.encryptedPayload; expiry timestamp supported.",
        "OAuth2: SharePoint client credentials flow (token cached + refresh logic).",
        "Browser auth: Playwright with TOTP for Jira SSO; 8-hour cookie session cache, file-backed.",
        "User authentication: NOT enforced. There is no login screen and no JWT validation; backend trusts every request.",
        "Multi-tenancy: hardcoded organization UUID; no isolation boundary — Phase 3 work.",
        "Audit: app.audit_log table exists but no middleware writes to it — Phase 2 work.",
        "CORS: open to all origins (development setting); needs locking down before any external exposure.",
    ])
    doc.add_page_break()


# ============== TARGET STATE ==============

def target_state():
    add_h1("6. Target State — Conceptual Architecture")
    add_body(
        "The platform is organised into nine functional modules. Each module is independently deployable and "
        "testable. The architecture is tech-stack-agnostic at the requirements level, but the implementation "
        "stack (React 19 + Express 5 + PostgreSQL + Redis + Playwright) is now locked-in for Phase 2 and beyond."
    )

    add_h2("6.1 System Architecture")
    add_body(
        "The diagram below shows how the four logical tiers — external systems, frontend, backend, and "
        "persistence — communicate. Source systems are on the orange row at top-left; destination systems "
        "are on the teal row at top-right. The React frontend talks only to the backend via the Fetch wrapper "
        "in services/api.js. The Express backend exposes 8 route groups, delegates to a service layer, and "
        "delegates long-running work to BullMQ workers backed by Redis. PostgreSQL holds all state."
    )
    add_image("architecture.png",
              caption="Figure 2 — High-level system architecture across the four tiers.",
              width_in=6.5)

    add_h2("6.2 Module Communication (Request Flow)")
    add_body(
        "A typical Operator-triggered run, traced as a sequence: the React page calls api.js, which posts to "
        "the Express route. The route validates input with Zod, calls a service, which writes a 'pending' "
        "row to runs and enqueues a BullMQ job. The route returns 202 Accepted immediately. The worker picks "
        "up the job asynchronously, executes the extract → normalise → map → push pipeline, updates run "
        "status, and the page sees the result via SSE or polling."
    )
    add_image("module_flow.png",
              caption="Figure 3 — Request flow for an Operator-initiated run, including async worker handoff.",
              width_in=6.5)

    add_h2("6.3 Module Catalog")
    add_table(
        ["Module", "Purpose", "Status Today"],
        [
            ["Connector Studio", "Design-time creation and configuration of connector templates", "Stub UI"],
            ["Connection Wizard", "Runtime guided setup of source-to-destination integrations", "Step 2 wired; rest stub"],
            ["Mapping Engine", "Visual and AI-assisted field mapping between entities", "Manual UI; AI missing"],
            ["Integration Registry", "Catalog of all published adapters and their instances", "Stub UI on mock data"],
            ["Health & Monitoring", "Real-time observability of all running integrations", "Stub UI on mock data"],
            ["Administration", "Platform governance — users, credentials, master data, approvals", "Stub UI; no auth backend"],
            ["Alerting & Notifications", "Proactive incident communication to stakeholders", "Missing"],
            ["Help System", "In-app guidance and documentation for all user personas", "HelpPanel widget only"],
            ["Web Scraping Engine", "Fallback data extraction when APIs are unavailable", "Missing"],
        ],
        col_widths_in=[2.0, 3.2, 1.3],
    )
    doc.add_page_break()


# ============== FUNCTIONAL REQUIREMENTS ==============

FR_TEMPLATE_NOTE = (
    "Each subsection below uses the same template: Description / User Stories / Acceptance Criteria / "
    "Backend impact / Frontend impact / DB / Dependencies / Phase."
)


def fr_block(num, title, description, stories, ac, backend, frontend, db, deps, phase):
    add_h2(f"7.{num} {title}")
    add_h3("Description")
    add_body(description)
    add_h3("User Stories")
    add_bullets(stories)
    add_h3("Acceptance Criteria")
    add_bullets(ac)
    add_h3("Backend Impact")
    add_bullets(backend)
    add_h3("Frontend Impact")
    add_bullets(frontend)
    add_h3("Database")
    add_body(db)
    add_h3("Dependencies")
    add_body(deps)
    add_h3("Phase")
    add_body(phase, bold=True)


def functional_requirements():
    add_h1("7. Functional Requirements (Future Features)")
    add_body(FR_TEMPLATE_NOTE)

    fr_block(
        1, "Connector Studio (Designer Only)",
        "Design-time workspace where a Designer onboards a new external system into the platform. "
        "One-time activity per system type. Once published, Operators can create unlimited instances. "
        "Sub-flows: System Registration → Authentication Configuration → Operation Selection → Entity "
        "Modelling → Publish & Version.",
        [
            "As a Designer, I select a system category (REST API, Database, File Share, SaaS, Message Queue, Webhook) and provide an OpenAPI/Swagger spec or DB connection string.",
            "As a Designer, I can group raw fields into entities with human-friendly labels and link them to master entities.",
            "As a Designer, I can publish a versioned connector template that Operators see in the Integration Registry.",
        ],
        [
            "Selecting REST API + uploading an OpenAPI 3.0 spec discovers all endpoints and lists them as selectable operations.",
            "Selecting Database + providing a Postgres connection string introspects the schema and lists tables.",
            "Each operation can be tagged Read / Write / Both and restricted from Operator visibility.",
            "Connector cannot be published until at least one successful test call with sample credentials.",
            "Updating a connector creates a new version; existing instances stay on prior version until explicitly upgraded.",
        ],
        [
            "New routes: POST /api/connectors, GET /api/connectors, POST /api/connectors/:id/discover, POST /api/connectors/:id/test, POST /api/connectors/:id/publish.",
            "OpenAPI parser service (use 'openapi-types' or 'swagger-parser').",
            "Database introspection service (Drizzle Kit's introspect or a custom pg metadata reader).",
            "Versioning logic: clone-on-publish, immutable past versions.",
        ],
        [
            "Studio page: enable Test Connection + Publish buttons (currently disabled at StudioPage.jsx:86).",
            "Sub-screens: System Registration form, OpenAPI uploader, schema introspection viewer, entity grouping drag-and-drop, publish confirmation modal.",
            "Spider/bot icon for web scraping connectors in registry.",
        ],
        "app.connectors becomes active. Add connector_versions table, connector_operations table, entity_definitions table.",
        "Phase 2 (Connector Studio backbone for new sources). OpenAPI parser, DB introspection libraries.",
        "Phase 2 — required before adding Dynamics 365, PostgreSQL, Excel, Keka, TARA connectors at scale.",
    )

    fr_block(
        2, "Connection Wizard (Operator)",
        "Guided 6-step flow with a visual stepper. Operator's primary workflow for creating an integration "
        "instance from published connectors.",
        [
            "As an Operator, I pick a Source and Destination from icon grids and toggle direction.",
            "As an Operator, I provide credentials with masked input and click Test Connection on each side.",
            "As an Operator, I select entities to map, then transition to Mapping Canvas.",
            "As an Operator, I configure schedule (cron-friendly UI), load strategy (full / incremental), retry policy.",
            "As an Operator, I run a dry-run with real sample data and see a split preview before clicking Publish.",
        ],
        [
            "Stepper shows Step X of 6 highlighted; cannot skip ahead.",
            "Test Connection per side independently with green ✓ / red ✗ result.",
            "Direction toggle: Source-to-Destination / Destination-to-Source / Bi-directional.",
            "Cron builder shows next-run preview in human language ('Every Mon 9:00 AM').",
            "Publish button disabled until dry-run passes — hard gate, no bypass.",
            "Confetti animation on successful Publish (charter §6.2 step 11).",
        ],
        [
            "Wizard step 4 currently redirects to Canvas — replace with inline mapping iframe / route preserving wizard state.",
            "Step 6 currently shows JSON preview — wire to actual dry-run endpoint POST /api/integrations/:id/dry-run.",
            "Persist wizard state to localStorage so refresh does not lose progress.",
        ],
        [
            "Backing endpoints already exist for Step 2; need /dry-run for Step 6.",
            "Integrate with Connector Studio output (entities, operations, credentials schema).",
        ],
        "Reuses app.integrations and runs tables — no new schema.",
        "Phase 2 — directly user-facing. Depends on Connector Studio (§7.1) and Mapping Canvas (§7.3).",
        "Phase 1 baseline → completed via Phase 2 wiring.",
    )

    fr_block(
        3, "Mapping Canvas (with AI Auto-Map)",
        "Full-screen drag-and-drop interface where source fields connect to destination fields via drawn "
        "lines. AI Auto-Map suggests mappings with confidence scores. Operator confirms, rejects, or "
        "adjusts. Transformation panel slides up from the bottom for selected lines.",
        [
            "As an Operator, I click AI Auto-Map and see dashed suggestion lines with confidence percentages.",
            "As an Operator, I confirm a suggestion (line goes solid green), reject it (line removed), or drag to a different target.",
            "As an Operator, I select preset transformations (date format, upper/lower, trim, concat, split) for any mapping.",
            "As an Operator, I describe a complex transformation in natural language and AI generates the JS.",
            "As an Operator, I can escalate an unmapped field as a task to the Designer with full context.",
        ],
        [
            "Lines colour-coded: green AI-confirmed, blue manual, amber pending review.",
            "AI mapping accuracy ≥70% in Phase 2 success metric.",
            "Validate button checks: all required dest fields mapped, types compatible, no circular dependencies.",
            "Mapping persists to integrations.fieldMappings JSONB on save.",
        ],
        [
            "New service: MappingAIService — wraps Claude API (anthropic SDK) for auto-map suggestions and NL→JS transforms.",
            "POST /api/integrations/:id/mappings/auto-map → returns suggestions with confidence scores.",
            "POST /api/integrations/:id/mappings/transform/nl → returns generated JS snippet.",
            "Sandbox JS evaluation for transformations (use vm2 or QuickJS) — never eval untrusted code on the main process.",
        ],
        [
            "Replace simulated AI in Canvas with real API calls.",
            "Connection lines via SVG with colour state machine.",
            "Slide-up transformation panel with preset dropdowns + Monaco/CodeMirror editor for advanced.",
            "Task escalation modal: pick Designer, attach source field + sample, send.",
        ],
        "No new tables — fieldMappings JSONB on app.integrations holds the structure.",
        "Anthropic Claude API key (env: ANTHROPIC_API_KEY). Sandbox library for JS eval.",
        "Phase 2 — high-impact UX feature.",
    )

    fr_block(
        4, "Integration Registry",
        "Searchable catalog of all published integration adapters and their running instances. Card-based "
        "layout with sparklines, status, and quick actions.",
        [
            "As an Operator, I search adapters by system type, department, status, or entity.",
            "As an Operator, I see message volume sparkline (last 7 days) and last run timestamp at a glance.",
            "As an Administrator, I bulk-pause or bulk-resume multiple adapters and export the registry as CSV.",
        ],
        [
            "Cards show source icon → destination icon, adapter name, status dot, sparkline, last run.",
            "Filters: system type, department, active/paused/error, entity involved.",
            "Detail expand shows: full config, mapping summary, schedule, recent run history, error log.",
            "Actions: Pause / Resume / Edit Mapping / Clone / Delete (with confirm) / View Logs.",
        ],
        [
            "Wire RegistryPage to GET /api/integrations + GET /api/runs?limit=7d for sparklines.",
            "Add POST /api/integrations/bulk-pause + bulk-resume.",
            "Add GET /api/integrations/export?format=csv.",
        ],
        [
            "Replace mock data import with services/api.js calls.",
            "Add charting library (Recharts recommended) for sparklines.",
            "Multi-select chip filtering already in mock UI — wire to query params.",
        ],
        "Reuses existing tables. Add view registry_view aggregating runs counts per integration.",
        "Charting library decision (Phase 2 Week 1).",
        "Phase 2 — primary Operator landing experience.",
    )

    fr_block(
        5, "Master Entity Catalog",
        "Central registry of all data entities defined across the platform, organised by department or use "
        "case (Engineering, HR, Finance, etc.). Supports merge with impact analysis and field-level "
        "evolution with cascade preview.",
        [
            "As a Designer, I see all entities in a Department > Entity tree.",
            "As a Designer, I select two overlapping entities, preview a side-by-side diff, and merge them — affected adapters are flagged for re-testing.",
            "As a Designer, I add a field to an entity and see which adapters need mapping updates before committing.",
        ],
        [
            "Tree navigation grouping by Department.",
            "Field-level usage statistics (used vs unused across adapters).",
            "Merge: side-by-side diff view, impact preview, post-merge re-test flag.",
            "Global search across all entities and fields.",
        ],
        [
            "New tables: entities, entity_fields, entity_field_versions.",
            "Endpoints: GET /api/entities (tree), POST /api/entities/:id/merge, POST /api/entities/:id/fields.",
            "Impact analysis service: query integrations.fieldMappings for usages.",
        ],
        [
            "Wire CatalogPage to GET /api/entities; replace inline hardcoded data.",
            "Tree component (use react-arborist or build from CSS).",
            "Side-by-side diff modal for merge preview.",
        ],
        "New tables (entities, entity_fields). Backfill from existing fieldMappings JSONB.",
        "Phase 2 (catalog read) + Phase 3 (merge / evolution).",
        "Phase 2 (read) → Phase 3 (write/merge).",
    )

    fr_block(
        6, "Trading Network Console (Message Monitor)",
        "Real-time and historical view of all data transactions flowing through the platform. Expandable rows "
        "show payload sent/received, field-level mapping trace, and transformation applied.",
        [
            "As an Administrator, I watch transactions in real-time with auto-refresh.",
            "As an Administrator, I expand an error row to see stack trace + which field caused failure.",
            "As an Administrator, I create a task for a Designer in one click from a failed row.",
            "As an Administrator, I export filtered results as CSV or JSON.",
        ],
        [
            "Table columns: Timestamp, Source, Destination, Adapter, Direction, Status, Records, Duration.",
            "Real-time toggle uses SSE (or polling every 5s).",
            "Filters: adapter, date range, status, source/dest, record count threshold.",
            "Error rows red-highlighted; expand reveals payload (JSON syntax-highlighted) + mapping trace.",
        ],
        [
            "Activate app.run_messages — write a row per message inside a run.",
            "Add SSE endpoint GET /api/messages/stream or polling endpoint GET /api/messages?since=.",
            "Add POST /api/tasks for one-click escalation.",
        ],
        [
            "Wire MonitorPage to live endpoint; replace mock monitorData.js.",
            "JSON viewer (react-json-view) with syntax highlighting.",
            "Pause button to freeze auto-refresh.",
        ],
        "app.run_messages becomes active. Add tasks table for escalation.",
        "Phase 2.",
        "Phase 2.",
    )

    fr_block(
        7, "Health Dashboard",
        "System-wide health at a glance. Landing page for Administrators, quick-check for Operators. KPI "
        "tiles, volume charts, client tier breakdown, error trend with anomaly detection.",
        [
            "As an Administrator, I see Total Adapters / Messages Today / Uptime / Active Alerts at the top.",
            "As an Administrator, I spot red adapter tiles in the grid (pulsing if unresolved).",
            "As an Administrator, I drill from any tile or chart to the relevant detail screen.",
        ],
        [
            "Time range selector (24h, 7d, 30d, custom).",
            "Volume chart: stacked bar (sent / received / errors) per adapter or aggregated.",
            "Client tier ribbon: Heavy / Moderate / Light usage.",
            "Error trend line with anomaly highlights.",
        ],
        [
            "Endpoints: GET /api/dashboard/summary, GET /api/dashboard/volume, GET /api/dashboard/errors-trend, GET /api/dashboard/client-tiers.",
            "Materialised views or scheduled aggregations to keep response time <500ms.",
        ],
        [
            "Wire DashboardPage to real endpoints; replace dashboardTiles.js mock.",
            "Charts: Recharts for volume bar, line for trend.",
            "Anomaly highlights: simple z-score per data point on the server.",
        ],
        "Add materialised view dashboard_volume_daily, dashboard_error_trend_hourly.",
        "Phase 2.",
        "Phase 2.",
    )

    fr_block(
        8, "Client & User Management + RBAC",
        "Registration and governance for external applications and internal users. Layered RBAC with four "
        "predefined roles (Administrator, Designer, Operator, Viewer) and granular per-resource overrides.",
        [
            "As an Administrator, I register a Consumer App and issue a client ID + secret.",
            "As an Administrator, I assign roles to users and override per-adapter permissions.",
            "As any user, I see my role badge and adapter ownership in the user directory.",
        ],
        [
            "RBAC matrix exactly as charter §5.8.2.",
            "Resource-level overrides: e.g., grant an Operator Designer access on one connector only.",
            "Action audit trail: append-only entries for publish, deploy, rotate, role change, alert ack.",
            "Session timeout: 30 min on vault, 8h general; force-logout for Administrators.",
        ],
        [
            "Auth middleware: JWT verify, decode user, load role, attach to req.user.",
            "Per-route guards using a withRole(['admin', 'designer']) decorator.",
            "Endpoints: POST /api/auth/login, POST /api/auth/refresh, GET /api/users, PATCH /api/users/:id/role, POST /api/clients/register.",
            "Activate app.audit_log via middleware on every state-changing route.",
        ],
        [
            "Login page (wire to /api/auth/login).",
            "AdminPage Users tab + Client Apps tab — replace mocks with real data.",
            "RBAC-aware nav: hide / disable nav items per role.",
            "Audit timeline view (filtered by user, entity, action).",
        ],
        "Activate app.users, app.organizations (real records), app.audit_log. Add app.user_overrides for resource-level grants.",
        "Choose JWT lib (jose recommended). Decide IdP (none for MVP — local accounts).",
        "Phase 2 (basic RBAC) → Phase 3 (resource overrides + audit trail UI).",
    )

    fr_block(
        9, "Credential Vault",
        "Centralised AES-256 storage with three modes: Runtime Prompt (never stored), Encrypted Vault, "
        "Browser Passthrough (OAuth). Eye-icon mask/reveal with audit logging on every reveal.",
        [
            "As an Operator, I store new credentials in the encrypted vault and see them masked by default.",
            "As an Operator, I click the eye icon to reveal a credential for 10s; the action is audit-logged.",
            "As an Administrator, I rotate or revoke credentials; expiry warnings appear at 7/3/1 days.",
        ],
        [
            "Reveal auto-hides after 10s or on blur.",
            "Copy-to-clipboard works without revealing.",
            "Sensitive values never appear in logs, exports, or error messages.",
            "Health report flags expiring + unused credentials.",
        ],
        [
            "Add GET /api/credentials/:id (returns metadata + decrypted value when authorised; logs the reveal).",
            "Rotate endpoint: PATCH /api/credentials/:id/rotate.",
            "Compliance report endpoint: GET /api/credentials/compliance.",
        ],
        [
            "VaultPage already has masked-reveal UI in mock — wire to real endpoints.",
            "Expiry countdown badge.",
            "Compliance report export (CSV) for audits.",
        ],
        "Add credential_reveals audit child of audit_log (or use audit_log with action='reveal').",
        "Encryption already implemented (AES-256-GCM). Need decrypt endpoint + audit hook.",
        "Phase 2.",
    )

    fr_block(
        10, "Help System",
        "Microsoft-style integrated help: searchable library, contextual '?' icon per screen, tooltips, "
        "first-time walkthroughs. AI Help Assistant scheduled for Phase 3+.",
        [
            "As any user, I open the right-side help panel and search 'how do I map a date field'.",
            "As a first-time user, I see a guided walkthrough overlay with numbered callouts on each major screen.",
            "As a Designer, I see inline info icons on complex form fields explaining expected formats.",
        ],
        [
            "Help panel slides in from right, pinnable.",
            "Search returns ranked results with highlighted matches.",
            "Tooltips on every icon and non-obvious UI element.",
            "First-visit walkthroughs are dismissable + re-accessible from Help menu.",
        ],
        [
            "Static help content as Markdown files in repo (charter §12 decision).",
            "Search index built at app build time (lunr.js or fuse.js).",
            "AI Assistant (Phase 3): wraps Claude API with screen state + recent actions context.",
        ],
        [
            "HelpPanel component already exists — wire to content engine.",
            "Walkthrough component (use shepherd.js or build with React portals).",
            "'?' icon on every page header, pre-filtered to current screen articles.",
        ],
        "No DB tables — Markdown in repo + build-time index.",
        "Markdown content authoring (Phase 1 Week 4 per charter open decision).",
        "Phase 2 (search + tooltips) → Phase 3 (AI Assistant).",
    )

    fr_block(
        11, "Alerting & Notifications (Break-Fix)",
        "Proactive incident communication. Triggers cover integration failure, degraded performance, "
        "credential expiry, schema change, scraping breakage. Multi-channel delivery, escalation chain, "
        "auto-pause on unacknowledged criticals.",
        [
            "As a Designer, I receive an in-app + email alert when my connector fails on a schema change.",
            "As an Operator, I see unresolved critical alerts as a banner at the top of every screen.",
            "As an Administrator, I configure custom thresholds (e.g., volume drops below 100 daily).",
        ],
        [
            "In-app bell with severity-sorted feed (critical > warning > info).",
            "Email: immediate critical, daily digest warnings, off info.",
            "Escalation: 30 min unacknowledged → admin; 2h → adapter auto-pause.",
            "Acknowledge / Create Task / View Logs from each alert detail.",
        ],
        [
            "Activate alert-dispatcher BullMQ worker.",
            "AlertService.create() called from run failure paths in integration-runner + sync workers.",
            "Email transport: nodemailer + SMTP env config (or SES / SendGrid).",
            "Endpoints: GET /api/alerts, POST /api/alerts/:id/acknowledge, POST /api/alerts/configure-thresholds.",
        ],
        [
            "AlertsPage already designed — wire to /api/alerts.",
            "Top-of-page banner component.",
            "Threshold config UI (per adapter or globally).",
        ],
        "Activate app.alerts. Add alert_thresholds, alert_subscriptions tables.",
        "SMTP credentials. nodemailer.",
        "Phase 2 (in-app + basic email) → escalation chains + thresholds.",
    )

    fr_block(
        12, "Web Scraping Engine",
        "Fallback data extraction when APIs are unavailable. Two sub-paths: Apify cloud actor and "
        "Playwright self-hosted automation. Output JSON entity matching API connector format so downstream "
        "mapping/monitoring is identical.",
        [
            "As a Designer, I select Web Scraping category and configure either Apify (target URL + selectors + schedule) or Playwright (recorded flow).",
            "As an Operator, I see scraping connectors flagged with a spider/bot icon in the registry.",
            "As a Designer, I receive a screenshot-diff alert when a scraping job fails on a UI change.",
        ],
        [
            "Both Apify and Playwright outputs match the entity schema produced by API connectors.",
            "Scheduled scrape jobs run on the same cron infrastructure.",
            "On scrape failure, alert includes a screenshot diff at the failure point.",
            "Anti-detection basics: rotating user agents, configurable delay between actions.",
        ],
        [
            "ApifyClient wrapper (uses official apify-client npm package).",
            "Playwright actor service (already have Playwright runtime for Flatiron).",
            "Worker: scrape-runner — separate concurrency limit (slow + IO-heavy).",
            "Failure path stores last screenshot for diff next run.",
        ],
        [
            "Studio: Scraping category sub-form (Apify config vs Playwright recorder).",
            "Spider/bot icon overlay on registry cards.",
        ],
        "Add scrape_jobs, scrape_runs, scrape_failures (with screenshot blob refs) tables.",
        "Apify account (env: APIFY_TOKEN). Object storage for screenshots (S3-compatible).",
        "Phase 2.",
    )
    doc.add_page_break()


# ============== NON-FUNCTIONAL ==============

def non_functional():
    add_h1("8. Non-Functional Requirements")
    add_h2("8.1 Security")
    add_bullets([
        "All credentials encrypted at rest with AES-256-GCM. Plaintext credentials never persisted, never logged.",
        "Eye-icon reveal events written to immutable audit log with user, timestamp, IP.",
        "RBAC enforced server-side on every state-changing route. Frontend role-gating is UX, not security.",
        "Session timeout: 30 min on vault screens, 8h on general screens. Force-logout for Administrators.",
        "CORS locked to known origins in production (currently open).",
        "Rate limiting at the Express layer (express-rate-limit) — 100 req/min per IP for unauthenticated, higher per-user post-auth.",
    ])
    add_h2("8.2 Performance")
    add_bullets([
        "Operator setup time < 30 minutes for an integration with pre-built connectors.",
        "Designer onboarding time < 2 hours for a new system type.",
        "Error detection latency < 5 minutes from failure to alert.",
        "Dashboard summary endpoint < 500 ms via materialised views.",
        "Mapping canvas remains interactive at ≥60 fps with 200+ field pairs.",
    ])
    add_h2("8.3 Reliability & Observability")
    add_bullets([
        "All BullMQ workers idempotent — re-running the same job must not produce duplicate SP items (already enforced by 3-layer dedup).",
        "Run history retained for 90 days; aggregated metrics retained for 1 year.",
        "Health endpoint /health and a dedicated /metrics (Prometheus format) for ops scraping.",
        "Structured logging (JSON lines) — pino recommended for backend.",
    ])
    add_h2("8.4 Accessibility & Browser Support")
    add_bullets([
        "Colour is never the only indicator (always paired with icon or text).",
        "Keyboard navigable (WCAG 2.1 AA target).",
        "Screen-reader compatible (semantic HTML, aria attributes).",
        "Browsers: Chrome, Edge, Firefox latest two versions.",
    ])
    add_h2("8.5 Scalability")
    add_bullets([
        "100 concurrent integrations per organisation with no degradation in dashboard responsiveness.",
        "10,000 messages/hour throughput at the worker level.",
        "Horizontal scale: add worker replicas without code change (BullMQ supports this natively).",
    ])
    doc.add_page_break()


# ============== COVERAGE MATRIX ==============

def coverage_matrix():
    add_h1("9. Source / Destination Coverage Matrix")
    add_body(
        "The screenshot accompanying this BRD shows the target system grid. Each row below maps that grid to "
        "an integration approach and the phase in which it will be built."
    )
    add_h2("9.1 Sources")
    add_table(
        ["System", "Approach", "Auth", "Status", "Phase"],
        [
            ["Jira", "REST API + Browser fallback", "Basic Auth + SSO/MFA via Playwright", "Built", "Done"],
            ["Dynamics 365", "REST API (Web API + OData)", "OAuth2 (Azure AD)", "Not started", "Phase 2"],
            ["TARA", "REST API (or scraping fallback)", "API Token", "Not started", "Phase 2"],
            ["PostgreSQL", "DB introspection + SQL", "Connection string + role", "Not started", "Phase 2"],
            ["Excel", "File parser (CSV / XLSX)", "Local / SharePoint / SFTP", "Not started", "Phase 2"],
            ["Keka", "REST API", "API Key", "Not started", "Phase 2"],
        ],
        col_widths_in=[1.2, 1.8, 1.5, 1.0, 1.0],
    )
    add_h2("9.2 Destinations")
    add_table(
        ["System", "Approach", "Auth", "Status", "Phase"],
        [
            ["SharePoint", "Microsoft Graph API", "OAuth2 client credentials", "Built", "Done"],
            ["TFS / Azure DevOps", "REST API", "Personal Access Token / OAuth", "Not started", "Phase 2"],
            ["Holiday Tracker", "REST API or DB write", "TBD per implementation", "Not started", "Phase 2"],
            ["Ahrefs", "REST API", "API Token", "Not started", "Phase 3"],
            ["Google Search Console", "REST API + service account", "Service Account / OAuth", "Not started", "Phase 3"],
            ["Google Adwords", "Google Ads API", "OAuth2 + developer token", "Not started", "Phase 3"],
        ],
        col_widths_in=[1.4, 1.7, 1.7, 0.9, 0.8],
    )
    doc.add_page_break()


# ============== ROADMAP ==============

def roadmap():
    add_h1("10. Phased Roadmap (Aggressive — Days)")
    add_body(
        "All durations are in working days for one full-time developer using Claude Code (vibe coding). The "
        "companion Synapse_Project_Plan_v1.xlsx breaks each item into a day-by-day sprint plan with goals, files "
        "touched, and tests-that-prove-done."
    )
    add_table(
        ["Phase", "Theme", "Approx Duration", "Headline Outcomes"],
        [
            ["Phase 1 (Done)", "Core pipeline (Jira → SharePoint)", "Already complete", "AES-256 vault, cron sync, 35-field mapper, delta sync, 3-layer dedup."],
            ["Phase 2", "Intelligence & Monitoring + multi-source", "≈ 35–45 days", "Studio backbone, RBAC, AI Auto-Map, Trading Console, Health Dashboard, Master Catalog (read), Web Scraping, alerts, Dynamics/PostgreSQL/Excel/Keka/TARA + TFS/Holiday Tracker connectors."],
            ["Phase 3", "Scale & Governance", "≈ 25–30 days", "Client App Registration, entity merge/evolution, self-healing, webhooks, bulk ops, audit log UI, AI Help Assistant, Ahrefs / GSC / Google Adwords destinations."],
            ["Phase 4", "Polish & Distribution", "≈ 15–20 days", "Theme engine, Mobile PWA, onboarding wizard, Docker packaging, full handover documentation."],
        ],
        col_widths_in=[1.2, 1.8, 1.3, 2.2],
    )
    doc.add_page_break()


# ============== RISKS & DECISIONS ==============

def risks_decisions():
    add_h1("11. Open Decisions & Risks")
    add_h2("11.1 Risk Register")
    add_table(
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["External API access not obtainable for some destinations", "Medium", "High", "Three-tier fallback: browser passthrough → Apify cloud → Playwright self-hosted. Scraping connectors flagged with bot icon."],
            ["Web scraping connectors break on UI changes", "High", "Medium", "Screenshot-diff alerting on scrape failures; designer notified immediately."],
            ["Scope creep expands beyond planned phases", "High", "High", "Lock scope per phase; change request process; MVP-first discipline."],
            ["AI mapper produces poor suggestions", "Medium", "Medium", "Always require human confirmation; log decisions for tuning; graceful fallback to manual mapping."],
            ["Non-technical users find mapping too complex", "Medium", "High", "Extensive UX testing; progressive disclosure; task escalation to Designer."],
            ["Credential vault security vulnerability", "Low", "Critical", "AES-256-GCM, security review before any external exposure, no raw credential display anywhere."],
            ["10 of 11 frontend pages on mock data — silent regressions when wired", "High", "Medium", "End-to-end Playwright test per page added at the time of wiring (test-first per page)."],
            ["No auth enforcement today — accidental public exposure if deployed", "Medium", "Critical", "Lock to localhost / VPN until auth ships in Phase 2 RBAC story."],
            ["Single-developer bus factor + vibe-coding velocity dependence", "Medium", "High", "Strong PRODUCT_STATUS + this BRD + day-level plan keep the project transferable; Claude session history saved with the repo."],
        ],
        col_widths_in=[2.4, 0.8, 0.8, 2.5],
    )
    add_h2("11.2 Open Decisions")
    add_table(
        ["Decision", "Options", "Recommendation", "Decide By"],
        [
            ["AI provider for mapping", "OpenAI / Claude API / Local model", "Claude API (already used elsewhere; cost-effective for the volume).", "Phase 2 Day 1"],
            ["Charting library", "Recharts / Chart.js / D3", "Recharts — works cleanly with React 19, declarative.", "Phase 2 Day 1"],
            ["Web scraping provider", "Apify (cloud) / Playwright (self) / Both", "Both — Apify for quick setups, Playwright for complex flows.", "Phase 2 Week 1"],
            ["Auth strategy", "Local (JWT) / SSO (Azure AD) / Both", "Local JWT first; pluggable adapter for SSO in Phase 3.", "Phase 2 Day 1"],
            ["Help content authoring", "Markdown in repo / CMS / Wiki", "Markdown in repo; index built at build-time.", "Phase 2 Week 4"],
            ["Email transport", "SMTP / SES / SendGrid", "Start SMTP via env; swap for SES in Phase 4.", "Phase 2 alerts kick-off"],
            ["JS sandbox for transformations", "vm2 / QuickJS / WASM", "QuickJS — modern, sandboxed, low overhead.", "Phase 2 mapping kick-off"],
        ],
        col_widths_in=[2.0, 2.0, 2.0, 1.0],
    )
    doc.add_page_break()


# ============== APPENDICES ==============

def appendices():
    add_h1("Appendix A — Folder Structure")
    add_h2("Frontend (packages/frontend/src)")
    add_table(
        ["Folder", "Purpose"],
        [
            ["components/admin", "User and client-app management UI"],
            ["components/alerts", "Alert cards, severity filtering, action buttons"],
            ["components/canvas", "Field mapping UI (pair source/dest fields)"],
            ["components/catalog", "Entity browser with tree navigation"],
            ["components/connected", "Integration dashboard, schedule and sync modals"],
            ["components/dashboard", "KPI tiles, status filters, time-range selector"],
            ["components/layout", "Topbar, Sidebar, DetailPane, HelpPanel, Toasts"],
            ["components/monitor", "Run history with expandable JSON payload viewer"],
            ["components/push", "Two-column form, integration selector, dedup modal"],
            ["components/registry", "Integration cards with sparklines, search, detail panes"],
            ["components/studio", "Connector builder shell (currently disabled buttons)"],
            ["components/vault", "Credential table with mask/reveal timer"],
            ["components/wizard", "6-step integration setup stepper"],
            ["contexts/", "React context providers (Theme, Toast, DetailPane, Sidebar)"],
            ["data/", "Mock data files (replace with real API as wiring progresses)"],
            ["hooks/", "Custom hooks (useDetailPane, etc.)"],
            ["services/api.js", "Centralised API client wrapper"],
        ],
        col_widths_in=[2.2, 4.3],
    )
    add_h2("Backend (packages/backend/src)")
    add_table(
        ["Folder", "Purpose"],
        [
            ["api/", "Express route handlers (8 route modules)"],
            ["db/schema.ts", "Drizzle table and enum definitions (14 tables)"],
            ["db/client.ts", "Database connection initialisation"],
            ["db/repositories/", "Data access objects per entity"],
            ["integrations/jira/", "Red Gold REST + Flatiron browser clients"],
            ["integrations/sharepoint/", "Type definitions + push service"],
            ["services/", "SyncService, SharePointPushService, PlaywrightAuthService, SchedulerService, MappingEngine, etc."],
            ["mappers/", "35-field Jira ↔ SharePoint mapping logic"],
            ["queues/", "BullMQ queue declarations"],
            ["workers/", "integration-runner, jira-sp-sync, playwright-sessions"],
            ["__tests__/", "Vitest test suite (e2e + unit)"],
        ],
        col_widths_in=[2.2, 4.3],
    )
    doc.add_page_break()

    add_h1("Appendix B — Glossary")
    add_table(
        ["Term", "Definition"],
        [
            ["Adapter", "A deployed instance of a connector with credentials, mapping, and schedule. Created by Operators."],
            ["Connector", "A reusable template that knows how to talk to a system type (Jira, SharePoint, Postgres). Created by Designers."],
            ["Entity", "A logical grouping of related fields (Project, Sprint, Issue). Lives in the Master Entity Catalog."],
            ["Designer", "Persona who builds connector templates and entities. Technical."],
            ["Operator", "Persona who deploys integrations using pre-built connectors. Non-technical."],
            ["Administrator", "Persona who governs the platform — users, credentials, approvals."],
            ["Run", "A single execution of an integration — has start, finish, status, record counts."],
            ["Push log", "History row written every time data is pushed to a destination — used for dedup."],
            ["Sync state", "Watermark row per integration tracking last successful sync."],
            ["Trading network", "Charter term for the live message-flow console (think: airport tower view of all transactions)."],
            ["Break-fix alert", "Notification fired when an integration fails or degrades."],
            ["Vibe coding", "AI-assisted development using Claude Code — the velocity model assumed by this plan."],
        ],
        col_widths_in=[1.5, 5.0],
    )
    doc.add_page_break()

    add_h1("Appendix C — Setup & Run")
    add_h2("Prerequisites")
    add_bullets([
        "Node.js 20+",
        "PostgreSQL 16 (local or via docker-compose)",
        "Redis 7 (local or via docker-compose)",
        "Playwright browsers installed (npx playwright install)",
    ])
    add_h2("First-Time Setup")
    add_bullets([
        "git clone the repo",
        "cp .env.example .env and fill in DATABASE_URL, REDIS_URL, ENCRYPTION_KEY (32-byte hex)",
        "docker-compose up -d (starts Postgres + Redis)",
        "npm install at the repo root (workspaces install both packages)",
        "Run drizzle migrations: cd packages/backend && npm run db:migrate",
    ])
    add_h2("Daily Run")
    add_bullets([
        "npm run dev — starts backend (port 4000) and frontend (port 5173) concurrently",
        "Visit http://localhost:5173",
        "Health probe: curl http://localhost:4000/health",
    ])
    doc.add_page_break()

    add_h1("Appendix D — Environment Variables")
    add_table(
        ["Variable", "Required", "Purpose"],
        [
            ["DATABASE_URL", "Yes", "Postgres connection string (e.g. postgres://localhost:5432/xnhub)"],
            ["REDIS_URL", "Yes", "Redis connection string for BullMQ"],
            ["ENCRYPTION_KEY", "Yes", "32-byte hex key for AES-256-GCM credential encryption"],
            ["PORT", "No (default 4000)", "Backend HTTP port"],
            ["JIRA_BASE_URL", "Per integration", "Provided at runtime; not a global env"],
            ["SHAREPOINT_TENANT_ID / CLIENT_ID / CLIENT_SECRET", "Per integration", "Stored in vault, not env"],
            ["ANTHROPIC_API_KEY", "Phase 2", "For AI Auto-Mapping and AI Help Assistant"],
            ["APIFY_TOKEN", "Phase 2", "For Apify scraping connectors"],
            ["SMTP_HOST / SMTP_USER / SMTP_PASS", "Phase 2", "Email alerts transport"],
            ["NODE_ENV", "Yes", "development | production"],
        ],
        col_widths_in=[2.6, 1.4, 2.5],
    )


# ============== BUILD ==============

setup_header_footer()
cover()
toc()
doc_control()
executive_summary()
vision()
personas()
current_state()
target_state()
functional_requirements()
non_functional()
coverage_matrix()
roadmap()
risks_decisions()
appendices()

doc.save(str(OUT))
print(f"WROTE: {OUT}")
print(f"SIZE:  {OUT.stat().st_size:,} bytes")
